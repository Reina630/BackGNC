"""
Script pour générer des documents de test pour le gestionnaire documentaire
"""
import os
from pathlib import Path
from datetime import datetime, timedelta
import random

# Pour générer des PDF
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("⚠️  reportlab non installé - les PDF seront des fichiers vides")
    print("   Installez avec: pip install reportlab")

# Pour générer des fichiers Office
try:
    from docx import Document
    from docx.shared import Inches
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("⚠️  python-docx non installé - les DOCX seront des fichiers vides")
    print("   Installez avec: pip install python-docx")

try:
    from openpyxl import Workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    print("⚠️  openpyxl non installé - les XLSX seront des fichiers vides")
    print("   Installez avec: pip install openpyxl")

# Pour générer des images
try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("⚠️  Pillow non installé - les images seront des fichiers vides")
    print("   Installez avec: pip install Pillow")


def create_test_folder(base_path="test_documents"):
    """Crée le dossier de test"""
    path = Path(base_path)
    path.mkdir(exist_ok=True)
    return path


def generate_pdf(filepath, title, content):
    """Génère un fichier PDF avec du contenu"""
    if REPORTLAB_AVAILABLE:
        c = canvas.Canvas(str(filepath), pagesize=A4)
        width, height = A4
        
        # Titre
        c.setFont("Helvetica-Bold", 24)
        c.drawString(50, height - 100, title)
        
        # Date
        c.setFont("Helvetica", 10)
        c.drawString(50, height - 130, f"Date: {datetime.now().strftime('%d/%m/%Y')}")
        
        # Contenu
        c.setFont("Helvetica", 12)
        y_position = height - 180
        for line in content.split('\n'):
            c.drawString(50, y_position, line)
            y_position -= 20
            if y_position < 100:
                c.showPage()
                y_position = height - 100
        
        c.save()
    else:
        # Création d'un fichier vide
        filepath.touch()


def generate_docx(filepath, title, content):
    """Génère un fichier DOCX avec du contenu"""
    if PYTHON_DOCX_AVAILABLE:
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
        doc.add_paragraph('')
        
        for paragraph in content.split('\n\n'):
            doc.add_paragraph(paragraph)
        
        doc.save(str(filepath))
    else:
        filepath.touch()


def generate_xlsx(filepath, title, data):
    """Génère un fichier Excel avec des données"""
    if OPENPYXL_AVAILABLE:
        wb = Workbook()
        ws = wb.active
        ws.title = title[:31]  # Excel limite à 31 caractères
        
        # Ajout des en-têtes
        for col_idx, header in enumerate(data['headers'], 1):
            ws.cell(row=1, column=col_idx, value=header)
        
        # Ajout des données
        for row_idx, row_data in enumerate(data['rows'], 2):
            for col_idx, value in enumerate(row_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        
        wb.save(str(filepath))
    else:
        filepath.touch()


def generate_image(filepath, title, size=(800, 600)):
    """Génère une image avec du texte"""
    if PIL_AVAILABLE:
        # Créer une image avec un fond coloré
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#FFA07A', '#98D8C8', '#F7DC6F']
        img = Image.new('RGB', size, color=random.choice(colors))
        
        draw = ImageDraw.Draw(img)
        
        # Essayer d'utiliser une police système
        try:
            font = ImageFont.truetype("arial.ttf", 40)
            small_font = ImageFont.truetype("arial.ttf", 20)
        except:
            font = ImageFont.load_default()
            small_font = ImageFont.load_default()
        
        # Ajouter le titre centré
        bbox = draw.textbbox((0, 0), title, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        x = (size[0] - text_width) / 2
        y = (size[1] - text_height) / 2
        
        draw.text((x, y), title, fill='white', font=font)
        
        # Ajouter la date
        date_text = datetime.now().strftime('%d/%m/%Y')
        date_bbox = draw.textbbox((0, 0), date_text, font=small_font)
        date_width = date_bbox[2] - date_bbox[0]
        draw.text(((size[0] - date_width) / 2, y + 60), date_text, fill='white', font=small_font)
        
        img.save(str(filepath))
    else:
        filepath.touch()


def main():
    """Fonction principale pour générer tous les fichiers de test"""
    print("🚀 Génération des documents de test...\n")
    
    # Créer le dossier de test
    test_folder = create_test_folder()
    print(f"📁 Dossier créé: {test_folder.absolute()}\n")
    
    # Liste des documents à générer
    documents = [
        {
            'type': 'pdf',
            'filename': 'Rapport_Financier_Q4_2024.pdf',
            'title': 'Rapport Financier Q4 2024',
            'content': '''Ce rapport présente les résultats financiers du quatrième trimestre 2024.
            
Résumé Exécutif:
- Chiffre d'affaires: 2.4M€
- Croissance: +15% vs Q4 2023
- Marge opérationnelle: 23%

Détails par département:
• Ventes: Performance exceptionnelle
• Marketing: ROI en amélioration
• Operations: Coûts optimisés

Recommandations pour Q1 2025:
1. Maintenir la dynamique commerciale
2. Investir dans l'innovation
3. Renforcer les équipes'''
        },
        {
            'type': 'docx',
            'filename': 'Contrat_Partenariat_ACME.docx',
            'title': 'Contrat de Partenariat ACME',
            'content': '''CONTRAT DE PARTENARIAT

Entre les soussignés:
- Société IMAN GED, représentée par son directeur général
- Société ACME Corporation, représentée par son président

Il a été convenu ce qui suit:

Article 1 - Objet du contrat
Le présent contrat a pour objet de définir les modalités de collaboration entre les deux parties dans le cadre du développement commercial.

Article 2 - Durée
Le contrat est conclu pour une durée de 24 mois à compter de la signature.

Article 3 - Engagements
Chaque partie s'engage à respecter ses obligations contractuelles et à collaborer de bonne foi.'''
        },
        {
            'type': 'pdf',
            'filename': 'Facture_2024-0891.pdf',
            'title': 'Facture #2024-0891',
            'content': '''FACTURE

Numéro: 2024-0891
Date: 10/01/2025
Échéance: 09/02/2025

Client:
ACME Corporation
123 Avenue des Champs
75008 Paris

Prestation:
- Licence logicielle annuelle: 5,000€
- Support premium: 1,200€
- Formation utilisateurs: 800€

Sous-total HT: 7,000€
TVA 20%: 1,400€
Total TTC: 8,400€

Conditions de paiement: 30 jours'''
        },
        {
            'type': 'xlsx',
            'filename': 'Budget_Previsionnel_2025.xlsx',
            'title': 'Budget 2025',
            'data': {
                'headers': ['Département', 'Q1', 'Q2', 'Q3', 'Q4', 'Total'],
                'rows': [
                    ['Ventes', 250000, 280000, 290000, 320000, 1140000],
                    ['Marketing', 45000, 50000, 48000, 55000, 198000],
                    ['R&D', 120000, 125000, 130000, 135000, 510000],
                    ['Operations', 80000, 82000, 85000, 88000, 335000],
                    ['RH', 35000, 35000, 38000, 40000, 148000],
                    ['Total', 530000, 572000, 591000, 638000, 2331000]
                ]
            }
        },
        {
            'type': 'pdf',
            'filename': 'Proces_verbal_AG_2024.pdf',
            'title': 'Procès-verbal AG 2024',
            'content': '''PROCÈS-VERBAL
Assemblée Générale Ordinaire 2024

Date: 8 janvier 2025
Lieu: Siège social

Présents:
- M. Amadou D., Président
- Mme Marie L., Directrice Financière
- M. Ibrahim K., Directeur Technique
- Mme Fatou S., Directrice RH

Ordre du jour:
1. Approbation des comptes 2024
2. Budget prévisionnel 2025
3. Stratégie de développement
4. Questions diverses

Résolutions adoptées:
• Approbation des comptes 2024 à l'unanimité
• Budget 2025 approuvé avec 1 abstention
• Plan stratégique validé

La séance est levée à 17h30.'''
        },
        {
            'type': 'image',
            'filename': 'Photo_Evenement_Lancement.jpg',
            'title': 'Lancement Produit',
        },
        {
            'type': 'pdf',
            'filename': 'Manuel_Utilisateur_v3.pdf',
            'title': 'Manuel Utilisateur v3.0',
            'content': '''MANUEL UTILISATEUR
Version 3.0

Table des matières:
1. Introduction
2. Installation
3. Prise en main
4. Fonctionnalités avancées
5. Dépannage

1. INTRODUCTION
Bienvenue dans IMAN GED v3.0, votre solution de gestion documentaire.

2. INSTALLATION
Configuration requise:
- Windows 10/11 ou macOS 10.15+
- 4 GB RAM minimum
- 500 MB d'espace disque

3. PRISE EN MAIN
Première connexion:
• Entrez vos identifiants
• Configurez votre profil
• Importez vos premiers documents

4. FONCTIONNALITÉS
• Gestion des documents
• Partage et collaboration
• Recherche avancée
• Versioning automatique'''
        },
        {
            'type': 'xlsx',
            'filename': 'Tableau_RH_Janvier.xlsx',
            'title': 'RH Janvier',
            'data': {
                'headers': ['Employé', 'Département', 'Statut', 'Jours Congés', 'Évaluation'],
                'rows': [
                    ['Amadou D.', 'Direction', 'CDI', 5, 'Excellent'],
                    ['Marie L.', 'Finance', 'CDI', 3, 'Très bien'],
                    ['Ibrahim K.', 'Technique', 'CDI', 7, 'Excellent'],
                    ['Fatou S.', 'RH', 'CDI', 2, 'Bien'],
                    ['Sophie M.', 'Marketing', 'CDD', 4, 'Très bien'],
                    ['Pierre L.', 'Ventes', 'CDI', 6, 'Bien'],
                ]
            }
        },
        {
            'type': 'image',
            'filename': 'Logo_Entreprise.png',
            'title': 'Logo IMAN GED',
        },
        {
            'type': 'docx',
            'filename': 'Note_Service_Securite.docx',
            'title': 'Note de Service - Sécurité',
            'content': '''NOTE DE SERVICE

Objet: Nouvelles procédures de sécurité informatique
Date: 12 février 2026

À l'attention de tous les collaborateurs,

Dans le cadre du renforcement de notre politique de sécurité, nous mettons en place de nouvelles procédures:

1. Mots de passe
• Changement obligatoire tous les 90 jours
• Minimum 12 caractères
• Authentification à deux facteurs

2. Documents sensibles
• Classification obligatoire
• Chiffrement pour les documents confidentiels
• Sauvegarde quotidienne automatique

3. Accès externes
• VPN obligatoire en télétravail
• Validation des appareils personnels

Ces mesures entrent en vigueur le 1er mars 2026.

Merci de votre collaboration.'''
        },
        {
            'type': 'pdf',
            'filename': 'Presentation_Commerciale.pdf',
            'title': 'Présentation Commerciale',
            'content': '''PRÉSENTATION COMMERCIALE
IMAN GED - Solutions Documentaires

QUI SOMMES-NOUS?
Leader dans la gestion documentaire depuis 2020
+500 clients satisfaits
Présence dans 15 pays

NOS SOLUTIONS:
✓ Gestion Électronique de Documents
✓ Archivage numérique sécurisé
✓ Workflow automatisé
✓ OCR et reconnaissance intelligente
✓ Collaboration en temps réel

AVANTAGES:
• Gain de temps: -60% sur la recherche
• Économies: -70% de papier
• Sécurité: Conformité RGPD
• Mobilité: Accès 24/7

TARIFS:
Starter: 29€/mois/utilisateur
Business: 49€/mois/utilisateur
Enterprise: Sur devis

Contact: commercial@imanged.com'''
        }
    ]
    
    # Générer chaque document
    created_count = 0
    for doc in documents:
        filepath = test_folder / doc['filename']
        
        try:
            if doc['type'] == 'pdf':
                generate_pdf(filepath, doc['title'], doc['content'])
                print(f"✅ PDF créé: {doc['filename']}")
            elif doc['type'] == 'docx':
                generate_docx(filepath, doc['title'], doc['content'])
                print(f"✅ DOCX créé: {doc['filename']}")
            elif doc['type'] == 'xlsx':
                generate_xlsx(filepath, doc['title'], doc['data'])
                print(f"✅ XLSX créé: {doc['filename']}")
            elif doc['type'] == 'image':
                generate_image(filepath, doc['title'])
                print(f"✅ Image créée: {doc['filename']}")
            
            created_count += 1
        except Exception as e:
            print(f"❌ Erreur pour {doc['filename']}: {e}")
    
    print(f"\n✨ {created_count}/{len(documents)} documents créés avec succès!")
    print(f"📂 Emplacement: {test_folder.absolute()}")
    
    # Suggestions
    print("\n💡 Suggestions:")
    print("   • Utilisez ces fichiers pour tester l'import dans votre GED")
    print("   • Testez les filtres par type de fichier")
    print("   • Vérifiez la recherche et les métadonnées")
    print("   • Testez le téléchargement et la prévisualisation")


if __name__ == "__main__":
    main()
